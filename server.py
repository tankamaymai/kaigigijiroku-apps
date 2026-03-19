#!/usr/bin/env python3
"""
議事録メーカー - Webサーバー
FastAPI + Whisper + ChatGPT / Gemini / Ollama
"""
import os
import json
import glob
import subprocess
import tempfile
import shutil
from datetime import datetime
from pathlib import Path
import threading
from queue import Queue, Empty
from typing import Any, Callable, Optional

import httpx
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from openpyxl import load_workbook

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from google import genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    from faster_whisper import WhisperModel
    FASTER_WHISPER_AVAILABLE = True
except ImportError:
    FASTER_WHISPER_AVAILABLE = False

# ========================================
# 設定
# ========================================

APP_DIR = Path(__file__).parent
TEMPLATE_DIR = APP_DIR / "templates"
EXCEL_TEMPLATE_DIR = APP_DIR / "excel_templates"
OUTPUT_DIR = APP_DIR / "output"
WEB_DIR = APP_DIR / "web"
DICTIONARY_FILE = APP_DIR / "dictionary.json"

# outputフォルダがなければ作成
OUTPUT_DIR.mkdir(exist_ok=True)

# ========================================
# FastAPI アプリ
# ========================================

app = FastAPI(
    title="議事録メーカー API",
    description="音声から議事録を自動生成するAPI",
    version="2.0.0"
)

# 静的ファイル配信
app.mount("/static", StaticFiles(directory=WEB_DIR / "static"), name="static")


# ========================================
# ユーティリティ関数
# ========================================

def list_templates() -> list:
    """テンプレート一覧を取得"""
    files = sorted(glob.glob(str(TEMPLATE_DIR / "*.json")))
    templates = []
    for f in files:
        with open(f, "r", encoding="utf-8") as fp:
            data = json.load(fp)
            templates.append({
                "filename": os.path.basename(f),
                "name": data.get("name", ""),
                "sections": data.get("sections", [])
            })
    return templates


def load_template(filename: str) -> dict:
    """テンプレートを読み込む"""
    path = TEMPLATE_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"テンプレートが見つかりません: {filename}")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_dictionary() -> dict[str, str]:
    """文字起こし辞書を読み込む。キー=読み/誤表記、値=正しい表記"""
    if not DICTIONARY_FILE.exists():
        return {}
    try:
        with open(DICTIONARY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            return {str(k): str(v) for k, v in data.items()}
    except (json.JSONDecodeError, IOError):
        return {}


def save_dictionary(entries: dict[str, str]) -> None:
    """文字起こし辞書を保存"""
    with open(DICTIONARY_FILE, "w", encoding="utf-8") as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)


def apply_dictionary(transcript: str, dictionary: dict[str, str]) -> str:
    """文字起こし結果に辞書を適用。長いキーから順に置換（部分一致を防ぐ）"""
    if not dictionary:
        return transcript
    result = transcript
    # 長いキーから順に置換（「福岡和白」が「福岡」より先にマッチするように）
    for key in sorted(dictionary.keys(), key=len, reverse=True):
        if key and key in result:
            result = result.replace(key, dictionary[key])
    return result


_whisper_model_cache: dict = {}


def get_audio_duration_seconds(audio_path: str) -> float:
    """ffprobe で音声の長さ（秒）を取得。失敗時は 0"""
    try:
        r = subprocess.run(
            [
                "ffprobe",
                "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                audio_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if r.returncode != 0 or not r.stdout.strip():
            return 0.0
        return float(r.stdout.strip())
    except (ValueError, subprocess.TimeoutExpired, FileNotFoundError, OSError):
        return 0.0


def run_whisper(
    audio_path: str,
    model_name: str = "medium",
    dictionary: Optional[dict[str, str]] = None,
    progress_callback: Optional[Callable[[dict[str, Any]], None]] = None,
) -> str:
    """Whisperで文字起こし（faster-whisper優先、フォールバックで従来版）。辞書があれば反映"""
    if dictionary is None:
        dictionary = load_dictionary()
    if FASTER_WHISPER_AVAILABLE:
        return _run_faster_whisper(audio_path, model_name, dictionary, progress_callback)
    return _run_whisper_cli(audio_path, model_name, dictionary, progress_callback)


def _run_faster_whisper(
    audio_path: str,
    model_name: str = "medium",
    dictionary: Optional[dict[str, str]] = None,
    progress_callback: Optional[Callable[[dict[str, Any]], None]] = None,
) -> str:
    """faster-whisperで高速文字起こし（従来の4〜8倍速）"""
    import time
    start = time.time()

    print(f"[WHISPER] faster-whisper で実行")
    print(f"[WHISPER] 入力ファイル: {audio_path}")
    print(f"[WHISPER] モデル: {model_name}")

    if model_name not in _whisper_model_cache:
        print(f"[WHISPER] モデル '{model_name}' を読み込み中...")
        _whisper_model_cache[model_name] = WhisperModel(
            model_name,
            device="cpu",
            compute_type="int8",
        )
        print(f"[WHISPER] モデル読み込み完了")

    model = _whisper_model_cache[model_name]

    # 辞書の正しい表記をhotwordsとして渡す（Whisperが認識しやすくする）
    hotwords_str = None
    if dictionary:
        hotwords_str = " ".join(dictionary.values())
        print(f"[WHISPER] 辞書適用: {len(dictionary)}件")

    duration_sec = get_audio_duration_seconds(audio_path)
    if progress_callback:
        progress_callback({
            "type": "transcribe_start",
            "duration": duration_sec,
            "message": "文字起こしを開始しました",
        })

    segments, info = model.transcribe(
        audio_path,
        language="ja",
        beam_size=5,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
        hotwords=hotwords_str,
    )

    texts = []
    for segment in segments:
        texts.append(segment.text.strip())
        if progress_callback:
            end_t = float(getattr(segment, "end", 0) or 0)
            preview = (segment.text or "").strip()[:120]
            if duration_sec > 0:
                pct = min(99.9, 100.0 * end_t / duration_sec)
            else:
                pct = min(99.9, len(texts) * 2.0)
            progress_callback({
                "type": "transcribe_segment",
                "seconds_start": float(getattr(segment, "start", 0) or 0),
                "seconds_end": end_t,
                "duration": duration_sec,
                "percent": round(pct, 1),
                "preview": preview,
                "segment_index": len(texts),
            })

    transcript = "\n".join(texts)
    transcript = apply_dictionary(transcript, dictionary or {})
    elapsed = time.time() - start
    print(f"[WHISPER] 文字起こし完了: {len(transcript)}文字 ({elapsed:.1f}秒)")
    print(f"[WHISPER] 検出言語: {info.language} (確率: {info.language_probability:.2f})")

    if not transcript:
        raise HTTPException(status_code=500, detail="文字起こし結果が空です")

    if progress_callback:
        progress_callback({
            "type": "transcribe_done",
            "duration": duration_sec,
            "chars": len(transcript),
        })

    return transcript


def _run_whisper_cli(
    audio_path: str,
    model_name: str = "medium",
    dictionary: Optional[dict[str, str]] = None,
    progress_callback: Optional[Callable[[dict[str, Any]], None]] = None,
) -> str:
    """従来のWhisper CLIで文字起こし（フォールバック用）。辞書は後処理で適用"""
    print(f"[WHISPER] CLI版 (openai-whisper) で実行")
    print(f"[WHISPER] 入力ファイル: {audio_path}")
    print(f"[WHISPER] モデル: {model_name}")
    print(f"[WHISPER] 出力先: {OUTPUT_DIR}")

    duration_sec = get_audio_duration_seconds(audio_path)
    if progress_callback:
        progress_callback({
            "type": "transcribe_start",
            "duration": duration_sec,
            "message": "Whisper CLI で文字起こし中（進捗は推定できません。完了までお待ちください）",
        })
        progress_callback({
            "type": "transcribe_segment",
            "seconds_start": 0,
            "seconds_end": 0,
            "duration": duration_sec,
            "percent": 50.0,
            "preview": "",
            "segment_index": 0,
            "cli_mode": True,
        })

    result = subprocess.run(
        ["whisper", audio_path, "--language", "Japanese", "--model", model_name, "--output_dir", str(OUTPUT_DIR)],
        capture_output=True,
        text=True
    )

    print(f"[WHISPER] 終了コード: {result.returncode}")
    if result.stdout:
        print(f"[WHISPER] stdout: {result.stdout[:500]}")
    if result.stderr:
        print(f"[WHISPER] stderr: {result.stderr[:500]}")

    if result.returncode != 0:
        raise HTTPException(status_code=500, detail=f"Whisper実行エラー: {result.stderr}")

    base = os.path.splitext(os.path.basename(audio_path))[0]
    txt_path = OUTPUT_DIR / f"{base}.txt"

    if not txt_path.exists():
        txt_files = list(OUTPUT_DIR.glob("*.txt"))
        if txt_files:
            txt_path = max(txt_files, key=lambda p: p.stat().st_mtime)
        else:
            raise HTTPException(status_code=500, detail="文字起こしファイルが生成されませんでした")

    with open(txt_path, "r", encoding="utf-8") as f:
        transcript = f.read().strip()
    transcript = apply_dictionary(transcript, dictionary or {})
    if progress_callback:
        progress_callback({
            "type": "transcribe_done",
            "duration": duration_sec,
            "chars": len(transcript),
        })
    return transcript


def build_prompt(tpl: dict, transcript: str) -> str:
    """ChatGPT用プロンプトを構築"""
    sections = "\n".join([f'- "{s["key"]}": {s["label"]}の報告内容' for s in tpl["sections"]])
    rules = "\n".join([f"- {r}" for r in tpl.get("chatgpt_prompt", {}).get("style_rules", [])])
    
    # 期待するJSONの例を作成
    example_json = {s["key"]: f"（{s['label']}の内容）" for s in tpl["sections"]}
    import json
    example_str = json.dumps(example_json, ensure_ascii=False, indent=2)
    
    return f"""あなたは会議議事録作成のプロです。以下の文字起こしから、各項目の内容を抽出してください。

【テンプレート名】
{tpl["name"]}

【抽出する項目とキー】
{sections}

【ルール】
{rules}
- 該当する内容がない場合は空文字 "" を設定

【出力形式（厳守）】
以下の形式のJSONのみを返してください。コードブロック（```）は絶対に使わないでください。

{example_str}

【会議の文字起こし】
{transcript}
"""


def build_freeform_prompt(transcript: str) -> str:
    """テンプレートなし用のプロンプトを構築"""
    return f"""あなたは会議議事録作成のプロです。以下の文字起こしから、適切な議事録を作成してください。

【ルール】
- 会議の内容を分析し、適切な項目に分けて整理してください
- 箇条書きで簡潔にまとめてください
- 重要な決定事項やアクションアイテムは明確に記載してください
- 雑談や重複は除外してください
- 該当する内容がない項目は省略してください

【出力形式（厳守）】
以下のJSON形式のみを返してください。コードブロック（```）は絶対に使わないでください。
項目の値は文字列で返してください。箇条書きは改行で区切ってください。

{{
  "meeting_title": "会議のタイトル（内容から推測）",
  "date_info": "日時に関する情報（言及があれば）",
  "attendees": "参加者（言及があれば）",
  "agenda": "議題・アジェンダ",
  "discussion": "議論・報告内容の要約",
  "decisions": "決定事項",
  "action_items": "アクションアイテム・TODO",
  "notes": "その他特記事項"
}}

【会議の文字起こし】
{transcript}
"""


FREEFORM_LABELS = {
    "meeting_title": "会議タイトル",
    "date_info": "日時",
    "attendees": "参加者",
    "agenda": "議題・アジェンダ",
    "discussion": "議論・報告内容",
    "decisions": "決定事項",
    "action_items": "アクションアイテム",
    "notes": "その他特記事項",
}


def write_freeform_excel(data: dict, output_name: str) -> Path:
    """テンプレートなしでExcelファイルを生成"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "議事録"

    header_font = Font(name="Yu Gothic", size=14, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
    label_font = Font(name="Yu Gothic", size=11, bold=True, color="2B579A")
    label_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    value_font = Font(name="Yu Gothic", size=11)
    thin_border = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 80

    title = data.get("meeting_title", "会議議事録")
    ws.merge_cells("A1:B1")
    cell = ws["A1"]
    cell.value = title
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40

    row = 3
    for key, label in FREEFORM_LABELS.items():
        value = data.get(key, "")
        if not value or key == "meeting_title":
            continue

        if not isinstance(value, str):
            value = str(value)

        label_cell = ws.cell(row=row, column=1, value=label)
        label_cell.font = label_font
        label_cell.fill = label_fill
        label_cell.alignment = Alignment(vertical="top", wrap_text=True)
        label_cell.border = thin_border

        value_cell = ws.cell(row=row, column=2, value=value)
        value_cell.font = value_font
        value_cell.alignment = Alignment(vertical="top", wrap_text=True)
        value_cell.border = thin_border

        line_count = value.count("\n") + 1
        ws.row_dimensions[row].height = max(30, line_count * 18)

        row += 1

    output_path = OUTPUT_DIR / output_name
    wb.save(output_path)
    print(f"[EXCEL] 自由形式Excel生成完了: {output_path}")
    return output_path


def parse_ai_response(content: str, provider_name: str) -> dict:
    """AI応答からJSONをパースする共通処理"""
    if "```json" in content:
        content = content.split("```json")[1]
    if "```" in content:
        content = content.split("```")[0]
    content = content.strip()
    
    try:
        parsed = json.loads(content)
        print(f"[DEBUG] Parsed JSON keys: {list(parsed.keys())}")
        
        if "sections" in parsed and isinstance(parsed["sections"], dict):
            parsed = parsed["sections"]
            print(f"[DEBUG] Extracted from 'sections': {list(parsed.keys())}")
        
        return parsed
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"{provider_name}の応答をJSONとして解析できません: {e}\n応答: {content[:200]}")


def call_chatgpt(api_key: str, prompt: str, model: str = "gpt-4o-mini") -> dict:
    """ChatGPT APIを呼び出し"""
    if not OPENAI_AVAILABLE:
        raise HTTPException(status_code=500, detail="OpenAIライブラリがインストールされていません")
    
    client = OpenAI(api_key=api_key)
    
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "あなたは議事録作成のプロフェッショナルです。会議の内容を正確かつ簡潔にまとめてください。"},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=4000
    )
    
    content = response.choices[0].message.content
    print(f"[DEBUG] ChatGPT raw response: {content[:500]}...")
    
    return parse_ai_response(content, "ChatGPT")


def call_gemini(api_key: str, prompt: str, model: str = "gemini-2.0-flash") -> dict:
    """Gemini APIを呼び出し"""
    if not GEMINI_AVAILABLE:
        raise HTTPException(status_code=500, detail="Google GenAIライブラリがインストールされていません。pip install google-genai を実行してください。")
    
    client = genai.Client(api_key=api_key)
    
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=genai_types.GenerateContentConfig(
            system_instruction="あなたは議事録作成のプロフェッショナルです。会議の内容を正確かつ簡潔にまとめてください。",
            temperature=0.3,
            max_output_tokens=4000,
        )
    )
    
    content = response.text
    print(f"[DEBUG] Gemini raw response: {content[:500]}...")
    
    return parse_ai_response(content, "Gemini")


def call_ollama(prompt: str, model: str = "gemma3", endpoint: str = "http://localhost:11434") -> dict:
    """Ollama（ローカルLLM）を呼び出し。データは外部に送信されない。"""
    api_url = f"{endpoint.rstrip('/')}/api/chat"

    try:
        with httpx.Client(timeout=300.0) as client:
            response = client.post(api_url, json={
                "model": model,
                "messages": [
                    {"role": "system", "content": "あなたは議事録作成のプロフェッショナルです。会議の内容を正確かつ簡潔にまとめてください。"},
                    {"role": "user", "content": prompt}
                ],
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_predict": 4000,
                }
            })
            response.raise_for_status()
    except httpx.ConnectError:
        raise HTTPException(status_code=500, detail="Ollamaに接続できません。Ollamaが起動しているか確認してください。\n起動コマンド: ollama serve")
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=500, detail=f"Ollamaエラー: {e.response.text}")

    data = response.json()
    content = data.get("message", {}).get("content", "")
    print(f"[DEBUG] Ollama raw response: {content[:500]}...")

    return parse_ai_response(content, "Ollama")


def write_excel(tpl: dict, data: dict, output_name: str) -> Path:
    """Excelファイルを生成"""
    from openpyxl.styles import Alignment
    
    xlsx_path = APP_DIR / tpl["excel_template"]
    
    if not xlsx_path.exists():
        raise HTTPException(status_code=404, detail=f"Excelテンプレートが見つかりません: {xlsx_path}")
    
    wb = load_workbook(xlsx_path)
    ws = wb[tpl["sheet"]]
    
    print(f"[EXCEL] === 書き込み開始 ===")
    print(f"[EXCEL] テンプレート: {xlsx_path}")
    print(f"[EXCEL] データキー: {list(data.keys())}")
    print(f"[EXCEL] セクション: {[(s['key'], s['cell']) for s in tpl['sections']]}")
    
    written_count = 0
    for s in tpl["sections"]:
        key = s["key"]
        cell_ref = s["cell"]
        prefix = s.get("prefix", "")  # ラベルのプレフィックス
        value = data.get(key, "")
        
        # 値が文字列でない場合は文字列に変換
        if value is not None and not isinstance(value, str):
            value = str(value)
        
        if not value:
            print(f"[EXCEL] スキップ: '{key}' (値なし)")
            continue
        
        # プレフィックス（ラベル）+ データを結合して書き込む
        full_value = prefix + value if prefix else value
        ws[cell_ref] = full_value
        
        # テキスト折り返しと上揃えを設定
        ws[cell_ref].alignment = Alignment(wrap_text=True, vertical='top')
        
        print(f"[EXCEL] ✅ {cell_ref} ← '{key}': {full_value[:50]}...")
        written_count += 1
    
    print(f"[EXCEL] === 書き込み完了: {written_count}件 ===")
    
    output_path = OUTPUT_DIR / output_name
    wb.save(output_path)
    
    return output_path


def _data_to_labeled_items(data: dict, tpl: Optional[dict], is_freeform: bool) -> list[tuple[str, str]]:
    """AIデータを (ラベル, 値) のリストに変換（テキスト/Word出力用）"""
    items = []
    if is_freeform:
        for key, label in FREEFORM_LABELS.items():
            value = data.get(key, "")
            if value and key != "meeting_title":
                items.append((label, str(value) if not isinstance(value, str) else value))
    else:
        key_to_label = {s["key"]: s.get("label", s["key"]) for s in tpl["sections"]}
        for key, value in data.items():
            if value and key in key_to_label:
                items.append((key_to_label[key], str(value) if not isinstance(value, str) else value))
    return items


def write_text(data: dict, tpl: Optional[dict], is_freeform: bool, output_name: str) -> Path:
    """テキストファイルを生成"""
    title = (tpl["name"] if tpl else None) or data.get("meeting_title", "議事録")
    items = _data_to_labeled_items(data, tpl, is_freeform)
    
    lines = [f"# {title}", "", "=" * 40, ""]
    for label, value in items:
        lines.append(f"## {label}")
        lines.append(value)
        lines.append("")
    
    output_path = OUTPUT_DIR / output_name
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    
    print(f"[TEXT] テキストファイル生成完了: {output_path}")
    return output_path


def write_docx(data: dict, tpl: Optional[dict], is_freeform: bool, output_name: str) -> Path:
    """Wordファイルを生成"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx がインストールされていません。pip install python-docx を実行してください。")
    
    doc = Document()
    title = (tpl["name"] if tpl else None) or data.get("meeting_title", "議事録")
    items = _data_to_labeled_items(data, tpl, is_freeform)
    
    # タイトル
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for label, value in items:
        doc.add_heading(label, level=2)
        p = doc.add_paragraph(value)
        p.paragraph_format.space_after = Pt(12)
    
    output_path = OUTPUT_DIR / output_name
    doc.save(output_path)
    print(f"[DOCX] Wordファイル生成完了: {output_path}")
    return output_path


def execute_pipeline_sync(
    tmp_path: str,
    original_filename: str,
    whisper_model: str,
    gpt_model: str,
    template: str,
    api_key: Optional[str],
    ai_provider: str,
    ollama_endpoint: str,
    output_format: str,
    progress_callback: Optional[Callable[[dict[str, Any]], None]] = None,
) -> dict:
    """音声→文字起こし→AI→ファイル出力の一連処理（同期）。進捗は progress_callback で通知"""
    is_freeform = template == "__none__"

    print(f"[INFO] 音声ファイル: {tmp_path}")
    print(f"[INFO] モード: {'自由形式' if is_freeform else 'テンプレート'}")
    print(f"[INFO] AIプロバイダー: {ai_provider}")
    print(f"[INFO] 出力形式: {output_format}")

    print(f"[INFO] Whisper開始 (モデル: {whisper_model})")
    transcript = run_whisper(
        tmp_path,
        whisper_model,
        dictionary=None,
        progress_callback=progress_callback,
    )
    print(f"[INFO] 文字起こし完了: {len(transcript)}文字")

    tpl = None
    if is_freeform:
        print(f"[INFO] 自由形式モード: テンプレートなしで議事録作成")
        prompt = build_freeform_prompt(transcript)
    else:
        tpl = load_template(template)
        print(f"[INFO] テンプレート読み込み: {tpl['name']}")
        prompt = build_prompt(tpl, transcript)

    if progress_callback:
        progress_callback({"type": "ai_start", "message": "AIで要約・議事録を作成しています"})

    if ai_provider == "ollama":
        print(f"[INFO] Ollama呼び出し (モデル: {gpt_model}, エンドポイント: {ollama_endpoint})")
        data = call_ollama(prompt, gpt_model, ollama_endpoint)
        print(f"[INFO] Ollama応答受信（ローカル処理完了）")
    elif ai_provider == "gemini":
        print(f"[INFO] Gemini API呼び出し (モデル: {gpt_model})")
        data = call_gemini(api_key, prompt, gpt_model)
        print(f"[INFO] Gemini応答受信")
    else:
        print(f"[INFO] ChatGPT API呼び出し (モデル: {gpt_model})")
        data = call_chatgpt(api_key, prompt, gpt_model)
        print(f"[INFO] ChatGPT応答受信")

    if progress_callback:
        progress_callback({"type": "file_write_start", "message": "ファイルを書き出しています"})

    base = os.path.splitext(original_filename)[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    meeting_title = data.get("meeting_title", "議事録")
    tpl_name = tpl["name"] if tpl else meeting_title

    ext_map = {"excel": ".xlsx", "text": ".txt", "docx": ".docx"}
    ext = ext_map.get(output_format, ".xlsx")
    output_name = f"{tpl_name}_{base}_{timestamp}{ext}"

    if output_format == "excel":
        if is_freeform:
            output_path = write_freeform_excel(data, output_name)
        else:
            output_path = write_excel(tpl, data, output_name)
    elif output_format == "text":
        output_path = write_text(data, tpl, is_freeform, output_name)
    elif output_format == "docx":
        output_path = write_docx(data, tpl, is_freeform, output_name)
    else:
        output_format = "excel"
        output_name = f"{tpl_name}_{base}_{timestamp}.xlsx"
        if is_freeform:
            output_path = write_freeform_excel(data, output_name)
        else:
            output_path = write_excel(tpl, data, output_name)

    if is_freeform:
        summary_display = {
            FREEFORM_LABELS.get(k, k): v
            for k, v in data.items()
            if v
        }
    else:
        summary_display = data

    format_labels = {"excel": "Excel", "text": "テキスト", "docx": "Word"}
    print(f"[INFO] {format_labels.get(output_format, 'Excel')}生成完了: {output_path}")

    return {
        "success": True,
        "filename": output_name,
        "path": str(OUTPUT_DIR),
        "transcript_length": len(transcript),
        "transcript": transcript,
        "summary": summary_display,
        "sections": list(data.keys()),
    }


# ========================================
# APIエンドポイント
# ========================================

@app.get("/", response_class=HTMLResponse)
async def index():
    """メインページ"""
    index_path = WEB_DIR / "index.html"
    with open(index_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/template-manager", response_class=HTMLResponse)
async def template_manager():
    """テンプレート管理ページ"""
    page_path = WEB_DIR / "template-manager.html"
    with open(page_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/api/templates")
async def get_templates():
    """テンプレート一覧を取得"""
    return {"templates": list_templates()}


@app.get("/api/dictionary")
async def get_dictionary():
    """文字起こし辞書を取得"""
    return {"entries": load_dictionary()}


@app.post("/api/dictionary")
async def save_dictionary_api(body: dict = Body(...)):
    """文字起こし辞書を保存。body: {"entries": {"読み": "漢字", ...}}"""
    entries = body.get("entries", body)
    if not isinstance(entries, dict):
        raise HTTPException(status_code=400, detail="entries はオブジェクトである必要があります")
    save_dictionary(entries)
    return {"success": True, "count": len(entries)}


@app.post("/api/process")
async def process_audio(
    file: UploadFile = File(...),
    whisper_model: str = Form("medium"),
    gpt_model: str = Form("gpt-4o-mini"),
    template: str = Form("shozokucho.json"),
    api_key: Optional[str] = Form(None),
    ai_provider: str = Form("openai"),
    ollama_endpoint: str = Form("http://localhost:11434"),
    output_format: str = Form("excel"),
):
    """
    音声ファイルを処理して議事録を生成
    ai_provider: "openai", "gemini", "ollama"
    template: "__none__" の場合はテンプレートなし（自由形式）
    output_format: "excel", "text", "docx"
    """
    tmp_path = None

    if ai_provider != "ollama" and not api_key:
        provider_name = "Gemini" if ai_provider == "gemini" else "OpenAI"
        raise HTTPException(status_code=400, detail=f"{provider_name} API Keyが必要です")

    try:
        suffix = os.path.splitext(file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        print(f"[INFO] 音声ファイル保存: {tmp_path} ({len(content)} bytes)")

        return execute_pipeline_sync(
            tmp_path,
            file.filename,
            whisper_model,
            gpt_model,
            template,
            api_key,
            ai_provider,
            ollama_endpoint,
            output_format,
            progress_callback=None,
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.post("/api/process-stream")
async def process_audio_stream(
    file: UploadFile = File(...),
    whisper_model: str = Form("medium"),
    gpt_model: str = Form("gpt-4o-mini"),
    template: str = Form("shozokucho.json"),
    api_key: Optional[str] = Form(None),
    ai_provider: str = Form("openai"),
    ollama_endpoint: str = Form("http://localhost:11434"),
    output_format: str = Form("excel"),
):
    """
    処理中に NDJSON で進捗をストリーミング（文字起こし位置・プレビュー等）。
    最終行は type: done と result を含む。
    """
    if ai_provider != "ollama" and not api_key:
        provider_name = "Gemini" if ai_provider == "gemini" else "OpenAI"
        raise HTTPException(status_code=400, detail=f"{provider_name} API Keyが必要です")

    suffix = os.path.splitext(file.filename)[1]
    content = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    original_filename = file.filename or "audio"

    def ndjson_generator():
        q: Queue = Queue()
        sentinel = object()

        def worker():
            try:
                result = execute_pipeline_sync(
                    tmp_path,
                    original_filename,
                    whisper_model,
                    gpt_model,
                    template,
                    api_key,
                    ai_provider,
                    ollama_endpoint,
                    output_format,
                    progress_callback=lambda ev: q.put(ev),
                )
                q.put({"type": "done", "result": result})
            except HTTPException as he:
                detail = he.detail
                if not isinstance(detail, str):
                    detail = str(detail)
                q.put({"type": "error", "detail": detail})
            except Exception as e:
                print(f"[ERROR] stream pipeline: {type(e).__name__}: {e}")
                q.put({"type": "error", "detail": str(e)})
            finally:
                q.put(sentinel)
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except OSError:
                        pass

        threading.Thread(target=worker, daemon=True).start()

        while True:
            try:
                item = q.get(timeout=0.3)
            except Empty:
                yield (json.dumps({"type": "ping"}, ensure_ascii=False) + "\n").encode("utf-8")
                continue
            if item is sentinel:
                break
            yield (json.dumps(item, ensure_ascii=False) + "\n").encode("utf-8")

    return StreamingResponse(
        ndjson_generator(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


def _get_media_type(filename: str) -> str:
    """ファイル拡張子からメディアタイプを返す"""
    ext = os.path.splitext(filename)[1].lower()
    return {
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".txt": "text/plain; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }.get(ext, "application/octet-stream")


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """生成したファイルをダウンロード（Excel / テキスト / Word）"""
    file_path = OUTPUT_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="ファイルが見つかりません")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=_get_media_type(filename)
    )


@app.post("/api/open-folder")
async def open_folder():
    """出力フォルダを開く"""
    subprocess.run(["open", str(OUTPUT_DIR)])
    return {"success": True}


@app.get("/api/health")
async def health_check():
    """ヘルスチェック"""
    return {
        "status": "ok",
        "openai_available": OPENAI_AVAILABLE,
        "gemini_available": GEMINI_AVAILABLE,
        "faster_whisper": FASTER_WHISPER_AVAILABLE,
        "templates_count": len(list_templates())
    }


# ========================================
# Ollama API
# ========================================

@app.get("/api/ollama/status")
async def ollama_status(endpoint: str = "http://localhost:11434"):
    """Ollamaの接続状態を確認"""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{endpoint.rstrip('/')}/api/tags")
            resp.raise_for_status()
            data = resp.json()
            models = [m["name"] for m in data.get("models", [])]
            return {"running": True, "models": models}
    except Exception:
        return {"running": False, "models": []}


@app.get("/api/ollama/models")
async def ollama_models(endpoint: str = "http://localhost:11434"):
    """Ollamaで利用可能なモデル一覧を取得"""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{endpoint.rstrip('/')}/api/tags")
            resp.raise_for_status()
            data = resp.json()
            models = []
            for m in data.get("models", []):
                size_gb = m.get("size", 0) / (1024**3)
                models.append({
                    "name": m["name"],
                    "size": f"{size_gb:.1f}GB",
                    "modified": m.get("modified_at", ""),
                })
            return {"models": models}
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Ollamaに接続できません。ollama serve で起動してください。")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========================================
# テンプレート管理API
# ========================================

from template_manager import analyze_excel_template, create_template_config, save_template_config


@app.post("/api/templates/upload")
async def upload_template_excel(file: UploadFile = File(...)):
    """
    Excelテンプレートをアップロードして解析
    """
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Excelファイル(.xlsx, .xls)のみ対応しています")
    
    # Excelファイルを保存
    excel_path = EXCEL_TEMPLATE_DIR / file.filename
    
    content = await file.read()
    with open(excel_path, "wb") as f:
        f.write(content)
    
    print(f"[INFO] Excelテンプレート保存: {excel_path}")
    
    # 解析
    try:
        analysis = analyze_excel_template(excel_path)
        return {
            "success": True,
            "filename": file.filename,
            "analysis": analysis
        }
    except Exception as e:
        # エラー時はファイルを削除
        if excel_path.exists():
            excel_path.unlink()
        raise HTTPException(status_code=500, detail=f"Excel解析エラー: {e}")


@app.post("/api/templates/create")
async def create_template(
    excel_filename: str = Form(...),
    template_name: str = Form(...),
    sections: str = Form(...),  # JSON文字列
    style_rules: str = Form(None)  # JSON文字列（オプション）
):
    """
    テンプレート設定を作成
    """
    excel_path = EXCEL_TEMPLATE_DIR / excel_filename
    
    if not excel_path.exists():
        raise HTTPException(status_code=404, detail=f"Excelファイルが見つかりません: {excel_filename}")
    
    try:
        sections_list = json.loads(sections)
        rules_list = json.loads(style_rules) if style_rules else None
        
        config = create_template_config(
            excel_path=excel_path,
            template_name=template_name,
            sections=sections_list,
            style_rules=rules_list
        )
        
        # 設定ファイルを保存
        config_path = save_template_config(config, TEMPLATE_DIR)
        
        print(f"[INFO] テンプレート設定保存: {config_path}")
        
        return {
            "success": True,
            "config": config,
            "config_file": config_path.name
        }
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"JSONパースエラー: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/templates/{template_name}/preview")
async def preview_template(template_name: str):
    """
    テンプレートのプレビュー情報を取得
    """
    tpl = load_template(template_name)
    excel_path = APP_DIR / tpl["excel_template"]
    
    if not excel_path.exists():
        raise HTTPException(status_code=404, detail="Excelテンプレートが見つかりません")
    
    wb = load_workbook(excel_path)
    ws = wb[tpl["sheet"]]
    
    # セル情報を取得
    cells_info = []
    for section in tpl["sections"]:
        cell = ws[section["cell"]]
        cells_info.append({
            "key": section["key"],
            "label": section["label"],
            "cell": section["cell"],
            "current_value": str(cell.value) if cell.value else "",
            "row": cell.row,
            "col": cell.column
        })
    
    return {
        "name": tpl["name"],
        "sheet": tpl["sheet"],
        "sections": cells_info,
        "style_rules": tpl.get("chatgpt_prompt", {}).get("style_rules", [])
    }


@app.delete("/api/templates/{template_name}")
async def delete_template(template_name: str):
    """
    テンプレートを削除
    """
    template_path = TEMPLATE_DIR / template_name
    
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="テンプレートが見つかりません")
    
    # テンプレート設定を読み込んでExcelファイルも削除
    try:
        with open(template_path, "r", encoding="utf-8") as f:
            config = json.load(f)
        
        excel_path = APP_DIR / config.get("excel_template", "")
        
        # 設定ファイル削除
        template_path.unlink()
        
        # Excelファイル削除（他のテンプレートで使用されていない場合）
        if excel_path.exists():
            # 他のテンプレートで使用されているかチェック
            other_templates = [t for t in list_templates() if t != template_name]
            is_used = False
            for other in other_templates:
                other_config = load_template(other)
                if other_config.get("excel_template") == config.get("excel_template"):
                    is_used = True
                    break
            
            if not is_used:
                excel_path.unlink()
        
        return {"success": True, "deleted": template_name}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========================================
# 起動
# ========================================

if __name__ == "__main__":
    import uvicorn
    print("=" * 50)
    print("📝 議事録メーカー Web版")
    print("=" * 50)
    print(f"🌐 http://localhost:8000")
    print("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=8000)
